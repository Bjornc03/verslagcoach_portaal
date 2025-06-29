import streamlit as st
from openai import OpenAI
import tempfile
import os
import docx
import fitz  # PyMuPDF
import win32com.client as win32
from docx import Document

# OpenAI-client initialiseren
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def extract_text(file):
    ext = file.name.split('.')[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file)
    elif ext == "docx":
        return extract_text_from_docx(file)
    else:
        return None

def extract_text_from_pdf(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name

    text = ""
    with fitz.open(tmp_path) as doc:
        for page in doc:
            text += page.get_text()
    os.remove(tmp_path)
    return text

def extract_text_from_docx(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name

    doc = Document(tmp_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    os.remove(tmp_path)
    return text

def generate_feedback(text, onderwerp, niveau):
    prompt = f"""
<verslagcoach>
  <instellingen>
    <taal>Nederlands</taal>
    <stijl>Professioneel, helder, concreet</stijl>
    <niveau>{niveau}</niveau>
  </instellingen>
  
  <verslag>
    <onderwerp>{onderwerp}</onderwerp>
    <tekst>
{text}
    </tekst>
  </verslag>
  
  <feedbackverzoek>
    <structuur>ja</structuur>
    <inhoudelijke_diepte>ja</inhoudelijke_diepte>
    <logica_argumentatie>ja</logica_argumentatie>
    <taalgebruik>ja</taalgebruik>
    <bronvermelding_APA>ja</bronvermelding_APA>
  </feedbackverzoek>
</verslagcoach>
    """

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Je bent een ervaren verslagcoach die gestructureerde, professionele feedback geeft."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return response.choices[0].message.content

def genereer_gesplitste_feedback(tekst, onderwerp, niveau):
    woorden = tekst.split()
    blokgrootte = 4000  # veilige marge onder 8192 tokens
    feedbacks = []

    for i in range(0, len(woorden), blokgrootte):
        deel = " ".join(woorden[i:i + blokgrootte])
        deelnummer = i // blokgrootte + 1
        st.info(f"🔍 Deel {deelnummer} wordt verwerkt...")
        try:
            deel_feedback = generate_feedback(deel, onderwerp, niveau)
            feedbacks.append((deelnummer, deel_feedback))
        except Exception as e:
            st.error(f"❌ Fout bij deel {deelnummer}: {e}")

    return feedbacks

def save_feedback_as_docx(feedback_delen, student_name):
    doc = Document()
    doc.add_heading(f"Verslagfeedback – {student_name}", level=1)

    for deelnummer, feedback_text in feedback_delen:
        doc.add_heading(f"📑 Feedback Deel {deelnummer}", level=2)
        for line in feedback_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    temp_path = tempfile.mktemp(suffix=".docx")
    doc.save(temp_path)
    return temp_path

def send_email_with_feedback(email, naam, feedback_path):
    outlook = win32.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.To = email
    mail.Subject = "Feedback op je verslag"
    mail.Body = f"Beste {naam},\n\nIn de bijlage vind je de feedback op je verslag.\n\nMet vriendelijke groet,\nVerslagcoach AI"
    mail.Attachments.Add(feedback_path)
    mail.Send()

# --- Streamlit interface ---
st.set_page_config(page_title="Verslagcoach", page_icon="📝")
st.title("📄 Verslagcoach – Uploadportaal")
st.write("Upload hier je verslag en ontvang gestructureerde feedback per e-mail.")

with st.form("upload_form"):
    naam = st.text_input("Je naam")
    email = st.text_input("Je e-mailadres")
    onderwerp = st.text_input("Waar gaat je verslag over?")
    niveau = st.selectbox("Opleidingsniveau", ["MBO", "HBO", "Universitair"])
    file = st.file_uploader("Upload je verslag (.docx of .pdf)", type=["docx", "pdf"])
    submitted = st.form_submit_button("Verstuur")

if submitted:
    if not all([naam, email, onderwerp, niveau, file]):
        st.warning("❗ Vul alle velden in en upload een bestand.")
    else:
        with st.spinner("📤 Bestand verwerken..."):
            verslagtekst = extract_text(file)
            if not verslagtekst or verslagtekst.strip() == "":
                st.error("⚠️ Er kon geen tekst uit het bestand worden gehaald. Controleer of het verslag niet leeg is.")
            else:
                feedback_delen = genereer_gesplitste_feedback(verslagtekst, onderwerp, niveau)
                if feedback_delen:
                    feedback_path = save_feedback_as_docx(feedback_delen, naam)
                    try:
                        send_email_with_feedback(email, naam, feedback_path)
                        st.success("✅ Feedback is verstuurd naar je e-mailadres.")
                    except Exception as e:
                        st.error(f"📧 Fout bij verzenden van e-mail: {e}")
